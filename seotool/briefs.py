"""Génération de livrables : brief d'article, brief de LP, brief d'article sponsorisé.

Le principe : l'outil assemble d'abord un *dossier de données* (requêtes GSC,
pages existantes proches, plans Hn internes, jargon, cibles de maillage), puis
un LLM ne fait que la mise en forme éditoriale. Les chiffres viennent toujours
de la base, jamais du modèle.
"""
from __future__ import annotations

import json
import os
import re
import textwrap
from datetime import date
from io import BytesIO

from . import gsc as gsc_mod
from . import semantic
from .store import Store


def dossier(store: Store, keyword: str, period: str | None = None, top_pages: int = 5) -> dict:
    """Rassemble tout ce que la base sait autour d'un mot-clé."""
    period = period or gsc_mod.last_n_months(1)[0]

    related = store.df(
        """SELECT query, SUM(impressions) impr, SUM(clicks) clics, AVG(position) pos,
                  COUNT(DISTINCT page) n_pages
           FROM gsc WHERE period >= ? AND query LIKE ?
           GROUP BY query ORDER BY impr DESC LIMIT 200""",
        (gsc_mod.last_n_months(6)[0], f"%{keyword}%"),
    )
    ranking_pages = store.df(
        """SELECT page, SUM(impressions) impr, SUM(clicks) clics, AVG(position) pos
           FROM gsc WHERE period >= ? AND query LIKE ?
           GROUP BY page ORDER BY impr DESC LIMIT ?""",
        (gsc_mod.last_n_months(6)[0], f"%{keyword}%", top_pages),
    )

    # pages internes les plus proches sémantiquement du mot-clé
    close = _closest_pages(store, keyword, k=top_pages)
    plans = {u: outline(store, u) for u in close}

    lexicon = [r[0] for r in store.conn.execute(
        "SELECT term FROM lexicon ORDER BY tfidf DESC LIMIT 60")]

    link_targets = store.df(
        """SELECT p.url, p.title, m.pagerank FROM pages p
           JOIN graph_metrics m ON m.url=p.url
           WHERE p.status=200 AND (p.title LIKE ? OR p.text LIKE ?)
           ORDER BY m.pagerank DESC LIMIT 10""",
        (f"%{keyword}%", f"%{keyword}%"),
    )


    return {
        "keyword": keyword,
        "period": period,
        "requetes": related.to_dict("records"),
        "volume_estime_mensuel": int(related["impr"].sum()) if len(related) else 0,
        "pages_positionnees": ranking_pages.to_dict("records"),
        "pages_proches": close,
        "plans_internes": plans,
        "jargon": lexicon,
        "cibles_maillage": link_targets.to_dict("records"),
        "cannibalisation": related[related["n_pages"] > 1].head(15).to_dict("records"),
    }


def _closest_pages(store: Store, keyword: str, k: int = 5) -> list[str]:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity

    rows = store.indexable_pages()
    docs = [f"{r['title'] or ''} {r['h1'] or ''} {r['text'] or ''}" for r in rows]
    vec = TfidfVectorizer(stop_words=semantic.FR_STOPWORDS, ngram_range=(1, 2), min_df=1)
    X = vec.fit_transform(docs + [keyword])
    sims = cosine_similarity(X[-1], X[:-1]).ravel()
    return [rows[i]["url"] for i in sims.argsort()[::-1][:k]]


def outline(store: Store, url: str) -> list[str]:
    rows = store.conn.execute(
        "SELECT level, text FROM headings WHERE url=? ORDER BY position", (url,)
    ).fetchall()
    return [f"{'  ' * (r['level'] - 1)}H{r['level']} {r['text']}" for r in rows]


# --- rendu markdown -------------------------------------------------------
def brief_markdown(d: dict, kind: str = "article") -> str:
    top = d["requetes"][:30]
    lignes = "\n".join(
        f"| {r['query']} | {int(r['impr'])} | {int(r['clics'])} | {r['pos']:.1f} |" for r in top
    )
    plans = "\n\n".join(
        f"**{u}**\n```\n" + "\n".join(p[:25]) + "\n```" for u, p in d["plans_internes"].items()
    )
    maillage = "\n".join(
        f"- [{r['title'] or r['url']}]({r['url']}) — PR {r['pagerank']:.4f}"
        for r in d["cibles_maillage"]
    )
    canni = "\n".join(
        f"- `{r['query']}` servie par {int(r['n_pages'])} URLs ({int(r['impr'])} impressions)"
        for r in d["cannibalisation"]
    ) or "- Aucune détectée"

    return textwrap.dedent(f"""\
    # Brief {kind} — {d['keyword']}
    _Généré le {date.today().isoformat()} · données GSC sur 6 mois glissants_

    ## 1. Potentiel
    - Volume estimé (impressions cumulées 6 mois) : **{d['volume_estime_mensuel']:,}**
    - Requêtes du champ : **{len(d['requetes'])}**
    - Pages déjà positionnées : {len(d['pages_positionnees'])}

    ## 2. Requêtes à couvrir (triées par impressions)
    | Requête | Impressions | Clics | Position moy. |
    |---|---:|---:|---:|
    {lignes}

    ## 3. Cannibalisation à surveiller
    {canni}

    ## 4. Plans Hn des pages internes proches (à ne pas dupliquer)
    {plans}

    ## 5. Maillage interne à prévoir
    Liens entrants à poser depuis :
    {maillage}

    ## 6. Champ lexical / jargon du site
    {', '.join(d['jargon'][:40])}

    ## 7. Plan proposé
    _(à compléter — voir génération LLM)_
    """).replace("{:,}", "")


def markdown_to_docx(markdown: str, title: str | None = None) -> bytes:
    """Convertit un brief Markdown en DOCX éditable, sans service externe."""
    from docx import Document
    from docx.shared import Pt

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    if title:
        document.core_properties.title = title

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for cell, value in zip(table.rows[0].cells, headers):
                cell.text = _plain_markdown(value)
            for values in rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, values):
                    cell.text = _plain_markdown(value)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            document.add_heading(_plain_markdown(heading.group(2)), level=min(len(heading.group(1)), 4))
        elif re.match(r"^[-*]\s+", stripped):
            document.add_paragraph(_plain_markdown(re.sub(r"^[-*]\s+", "", stripped)), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", stripped):
            document.add_paragraph(_plain_markdown(re.sub(r"^\d+[.)]\s+", "", stripped)), style="List Number")
        elif stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            document.add_paragraph("\n".join(block), style="No Spacing")
        else:
            document.add_paragraph(_plain_markdown(stripped))
        i += 1

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _plain_markdown(value: str) -> str:
    value = re.sub(r"\[([^]]+)]\(([^)]+)\)", r"\1 (\2)", value)
    return re.sub(r"[*_`]+", "", value).strip()


# --- génération LLM -------------------------------------------------------
EXIGENCES = """
Contraintes transverses à intégrer au brief :
- ALIGNEMENT À LA DEMANDE : appuie le brief sur les requêtes GSC, les pages déjà positionnées,
  les plans Hn internes et le vocabulaire réellement présent sur le site. Ne prétends pas connaître
  la SERP ou les concurrents si aucune donnée externe n'est fournie.
- E-E-A-T : indique au rédacteur quels éléments de première main sont attendus (donnée interne,
  cas client, test, citation d'expert nommé), quelles affirmations exigeront une source datée,
  et qui doit signer la page. Ce sont des proxys d'évaluation humaine, pas un score algorithmique :
  formule-les comme des exigences éditoriales, pas comme des règles Google.
- GEO (citabilité par les moteurs génératifs) : réponse directe et autoportante de 25 à 90 mots
  en ouverture, chaque section compréhensible hors contexte, H2 formulés comme des questions
  réelles, faits chiffrés attribués et datés, au moins un format extractible (tableau comparatif,
  liste de critères ou FAQ balisée FAQPage). Ces critères reflètent l'état de l'art observé, pas
  une mécanique documentée par les moteurs.
"""

PROMPTS = {
    "article": "Rédige un brief d'article SEO complet : angle, intention de recherche dominante, "
               "persona, plan Hn détaillé (H1-H3) avec pour chaque section les requêtes à couvrir "
               "et le nombre de mots cible, entités/jargon à employer, maillage interne, "
               "schema.org à baliser, meta title (<60c) et meta description (<155c), CTA.",
    "lp": "Rédige un brief de landing page B2B : promesse, preuve, structure de blocs (hero, "
          "bénéfices, preuve sociale, objections, FAQ balisée FAQPage), requêtes par bloc, "
          "wording des CTA, meta title et description.",
    "sponso": "Rédige un brief d'article sponsorisé destiné à un média partenaire (netlinking) : "
              "angle éditorial crédible pour le média, plan, emplacement naturel du lien, "
              "3 propositions d'ancres (exacte, semi-optimisée, marque) avec la phrase d'insertion, "
              "et un message de prise de contact au média.",
}


def generate(d: dict, kind: str = "article", model: str = "claude-sonnet-4-6") -> str:
    """Appelle l'API Anthropic pour transformer le dossier en brief rédigé."""
    import anthropic

    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    payload = json.dumps({k: v for k, v in d.items() if k != "plans_internes"},
                         ensure_ascii=False, default=str)[:120_000]
    plans = json.dumps(d["plans_internes"], ensure_ascii=False)[:30_000]

    msg = client.messages.create(
        model=model,
        max_tokens=4000,
        system=(
            "Tu es consultant SEO senior en agence B2B. Tu produis des briefs actionnables "
            "pour des rédacteurs. Tu t'appuies EXCLUSIVEMENT sur les données fournies "
            "(requêtes, impressions, positions, plans existants, jargon du site). "
            "Tu n'inventes aucun chiffre. Tu écris en français, en markdown."
        ),
        messages=[{
            "role": "user",
            "content": f"{PROMPTS[kind]}\n{EXIGENCES}\n\nMot-clé pivot : {d['keyword']}\n\n"
                       f"Données GSC et site :\n{payload}\n\nPlans Hn existants :\n{plans}",
        }],
    )
    return "".join(b.text for b in msg.content if b.type == "text")


def monthly_report(store: Store, period: str | None = None) -> str:
    """Le point mensuel : opportunités, déclins, cannibalisation, pages à créer."""
    period = period or gsc_mod.last_n_months(1)[0]
    opp = gsc_mod.opportunities(store, period).head(40)
    gaps = gsc_mod.gap_pages(store, period).head(30)
    canni = gsc_mod.cannibalisation(store, period).head(20)
    tr = gsc_mod.trend(store, 3)

    decline = ""
    if len(tr):
        piv = tr.pivot_table(index="page", columns="period", values="clics", aggfunc="sum").fillna(0)
        if piv.shape[1] >= 2:
            piv["delta"] = piv.iloc[:, -1] - piv.iloc[:, 0]
            decline = piv.sort_values("delta").head(20).to_markdown()

    return "\n\n".join([
        f"# Point SEO mensuel — {period}",
        "## Optimisations prioritaires (positions 4-20, fort volume)",
        opp.to_markdown(index=False) if len(opp) else "_Aucune_",
        "## Pages à créer (requêtes sans page dédiée)",
        gaps.to_markdown(index=False) if len(gaps) else "_Aucune_",
        "## Cannibalisation",
        canni.to_markdown(index=False) if len(canni) else "_Aucune_",
        "## Pages en déclin (3 mois)",
        decline or "_Historique insuffisant_",
    ])
